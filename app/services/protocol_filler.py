import os
import shutil
from pathlib import Path
from typing import Dict, Any
from docx import Document
from app.config.logging_config import get_logger
import re

# Initialize logger
logger = get_logger(__name__)


def replace_placeholders(paragraphs, data: Dict[str, Any]) -> int:
    replacements = 0
    pattern = re.compile(r"\{\{(\w+)\}\}")
    for p in paragraphs:
        text = p.text
        matches = list(pattern.finditer(text))
        if not matches:
            continue
        p.clear()
        last_index = 0
        for match in matches:
            start, end = match.span()
            key = match.group(1)
            # Add text before marker
            if start > last_index:
                p.add_run(text[last_index:start])
            # Add replacement with bold and italic if key in data
            if key in data:
                ans = p.add_run(str(data[key]))
                ans.bold = True
                ans.italic = True
                replacements += 1
            else:
                # If key not in data, add original marker text
                p.add_run(match.group(0))
            last_index = end
        # Add remaining text after last marker
        if last_index < len(text):
            p.add_run(text[last_index:])
    return replacements

def replace_placeholders_in_tables(paragraphs, data: Dict[str, Any]) -> int:
    replacements = 0
    pattern = re.compile(r"\{\{(\w+)\}\}")
    for p in paragraphs:
        text = p.text
        matches = list(pattern.finditer(text))
        if not matches:
            continue
        p.clear()
        last_index = 0
        for match in matches:
            start, end = match.span()
            key = match.group(1)
            # Add text before marker
            if start > last_index:
                p.add_run(text[last_index:start])
            # Add replacement without bold and italic if key in data
            if key in data:
                p.add_run(str(data[key]))
                replacements += 1
            else:
                # If key not in data, add original marker text
                p.add_run(match.group(0))
            last_index = end
        # Add remaining text after last marker
        if last_index < len(text):
            p.add_run(text[last_index:])
    return replacements

def fill_protocol(template_path: str, output_path: str, data: Dict[str, Any]) -> bool:
    """
    Fill a Word template with provided data and save to output path.

    Args:
        template_path: Path to the Word template file
        output_path: Path where to save the filled document
        data: Dictionary with key-value pairs to replace in the template

    Returns:
        bool: True if successful, False otherwise
    """
    logger.info(f"Starting to fill protocol template: {template_path}")
    logger.debug(f"Output path: {output_path}")
    logger.debug(f"Data keys: {list(data.keys())}")

    try:
        # Validate template exists
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            return False

        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            logger.debug(f"Created output directory: {output_dir}")

        # Create a copy of the template
        logger.debug("Creating template copy")
        shutil.copy(template_path, output_path)

        # Load the document
        doc = Document(output_path)

        # Process document content
        logger.debug("Processing document content")
        total_replacements = 0

        # Replace in paragraphs
        para_replacements = replace_placeholders(doc.paragraphs, data)
        total_replacements += para_replacements
        logger.debug(f"Made {para_replacements} replacements in paragraphs")

        table_replacements = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_replacements += replace_placeholders_in_tables(cell.paragraphs, data)
        total_replacements += table_replacements
        logger.debug(f"Made {table_replacements} replacements in tables")
        
        if total_replacements == 0:
            logger.warning("No template markers were replaced in the document")
        else:
            logger.info(f"Successfully made {total_replacements} replacements in the document")
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Successfully saved filled document to: {output_path}")
        return True
        
    except PermissionError as e:
        logger.error(f"Permission error when accessing files: {e}")
        return False
    except Exception as e:
        logger.error(f"Error filling protocol template: {e}", exc_info=True)
        # Clean up partially created file if it exists
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
                logger.debug(f"Removed partially created file: {output_path}")
            except Exception as cleanup_error:
                logger.error(f"Failed to clean up partially created file: {cleanup_error}")
        return False